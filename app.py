import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement 
from docx.oxml.ns import qn        
from groq import Groq
import json
import io
import os

# --- 1. INITIAL SESSION STATE ---
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if 'qp_bytes' not in st.session_state:
    st.session_state.qp_bytes = None
    st.session_state.ak_bytes = None
    st.session_state.current_sub_code = ""

# --- 2. CONFIGURATION (Groq) ---
api_key = os.environ.get("GROQ_API_KEY")

if not api_key:
    try:
        api_key = st.secrets["GROQ_API_KEY"]
    except (FileNotFoundError, KeyError):
        api_key = None

if api_key:
    client = Groq(api_key=api_key)
else:
    if st.session_state.logged_in:
        st.error("Please add GROQ_API_KEY to Render Environment Variables or secrets.toml")
        st.stop()

GROQ_MODEL = "llama-3.3-70b-versatile"

# --- 3. UTILITY FUNCTIONS ---
def extract_text_from_pdf(uploaded_file):
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    text = "".join([page.get_text() for page in doc])
    return text

def get_llm_response(raw_text, config):
    # Retrieve the user-defined exact question counts
    total_a = config['total_a']
    num_mcq = config['num_mcq']
    num_tf = config['num_tf']
    num_fill = config['num_fill']

    prompt = f"""
    Act as an academic expert and exam controller. Based on this source material: {raw_text[:15000]}
    
    Create a professional bilingual examination paper:
    - Subject: {config['subject_name']} ({config['subject_code']})
    - Branch: {config['branch_name']}
    
    Section Requirements:
    1. Section A: Generate exactly {total_a} questions. 
       - SEQUENCING & FORMATTING (CRITICAL RULE):
         1. First, generate EXACTLY {num_mcq} MCQs. You MUST include 4 options (A, B, C, D) directly inside the question text.
         2. Next, generate EXACTLY {num_fill} Fill-in-the-blanks questions.
         3. Finally, generate EXACTLY {num_tf} True/False questions.
       - DO NOT mix question types. Keep them strictly in the order above.
       - Students will attempt {config['n_a']}. Difficulty: {config['diff_a']}.
       
    2. Section B: Generate exactly {config['total_b']} questions. 
       - Short answer types. Students attempt {config['n_b']}. Difficulty: {config['diff_b']}.
       
    3. Section C: Generate exactly {config['total_c']} questions. 
       - Long answer/descriptive types. Students attempt {config['n_c']}. Difficulty: {config['diff_c']}.
    
    CRITICAL RULES:
    - Every question ("q") must be bilingual: [English Question] / [Hindi Translation].
    - Every answer ("a") must be clear and accurate.
    - Return ONLY a valid JSON object.
    
    JSON Structure:
    {{
        "section_a": [{{"q": "Q Text / प्रश्न \\n A) ... B) ... C) ... D) ...", "a": "Ans Text"}}, ...],
        "section_b": [{{"q": "Q Text / प्रश्न", "a": "Ans Text"}}, ...],
        "section_c": [{{"q": "Q Text / प्रश्न", "a": "Ans Text"}}, ...]
    }}
    """
    
    chat_completion = client.chat.completions.create(
        messages=[
            {"role": "system", "content": "You are an expert exam generator that outputs only JSON."},
            {"role": "user", "content": prompt}
        ],
        model=GROQ_MODEL,
        response_format={"type": "json_object"}
    )
    return json.loads(chat_completion.choices[0].message.content)

def create_word_files(data, config):
    def add_header(doc, is_ans=False):
        p_mm = doc.add_paragraph()
        p_mm.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_mm.add_run(f"Μ.Μ.: {config['total_marks']}").bold = True
        
        title_text = f"Answer Key: {config['title']}" if is_ans else config['title']
        h = doc.add_heading(title_text, 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        details = [
            f"Branch Name: {config['branch_name']} ({config['branch_code']})",
            f"Semester: {config['sem']} | Subject: {config['subject_name']} ({config['subject_code']})",
            f"Time: {config['duration']}"
        ]
        for line in details:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("_" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_custom_footer(doc, sub_code):
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        
        p.text = f"{sub_code}\tPage "
        
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.extend([fldChar1, instrText, fldChar2, fldChar3])
        

    # --- 1. Generate Question Paper ---
    q_doc = Document()
    add_header(q_doc)
    
    sections = [("A / भाग - क", "section_a", config['n_a'], config['m_a']),
                ("B / भाग - ख", "section_b", config['n_b'], config['m_b']),
                ("C / भाग - ग", "section_c", config['n_c'], config['m_c'])]
    
    q_counter = 1
    
    for label, key, req_count, sec_marks in sections:
        h = q_doc.add_heading(f"SECTION-{label}", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        q_doc.add_paragraph(f"Note: Attempt any {req_count} questions. (Total Marks: {sec_marks})").italic = True
        
        for item in data[key]:
            q_doc.add_paragraph(f"Q{q_counter}. {item['q']}")
            q_counter += 1
            
    add_custom_footer(q_doc, config['subject_code'])

    # --- 2. Generate Answer Key ---
    a_doc = Document()
    add_header(a_doc, is_ans=True)
    
    a_counter = 1
    
    for label, key, _, _ in sections:
        h = a_doc.add_heading(f"Answers: SECTION-{label}", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for item in data[key]:
            p = a_doc.add_paragraph()
            p.add_run(f"Q{a_counter}: ").bold = True
            p.add_run(item['q'])
            a_doc.add_paragraph(f"Ans: {item['a']}")
            a_counter += 1
            
    add_custom_footer(a_doc, config['subject_code'])
    
    return q_doc, a_doc

# --- 4. LOGIN PAGE LOGIC ---
if not st.session_state.logged_in:
    st.set_page_config(page_title="Login | SLOG AI", page_icon="🔐", layout="centered")
    
    st.markdown("""
        <style>
        #MainMenu, footer, header {visibility: hidden;}

        .stApp {
            background-image: url("https://images.unsplash.com/photo-1507842217343-583bb7270b66?q=80&w=2400&auto=format&fit=crop") !important;
            background-size: cover !important;
            background-position: center !important;
            background-attachment: fixed !important;
        }

        .main .block-container {
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            height: 100vh;
            padding: 0 !important;
        }

        div[data-testid="stForm"] {
            background-color: rgba(255, 255, 255, 0.95) !important;
            border-radius: 20px !important;
            box-shadow: 0 25px 50px -12px rgba(0, 0, 0, 0.5) !important;
            padding: 50px 40px !important;
            width: 100% !important;
            max-width: 400px !important;
            margin: auto !important;
            border: none !important;
            backdrop-filter: blur(10px);
        }

        .main-title {
            text-align: center;
            font-size: 2.2rem;
            font-weight: 800;
            color: #4c1d95; 
            margin-bottom: 2.5rem;
            font-family: 'Segoe UI', system-ui, sans-serif;
        }

        .stTextInput input {
            border: 2px solid #f1f5f9 !important;
            border-radius: 30px !important;
            padding: 16px 25px !important;
            font-size: 1rem !important;
            font-weight: 500 !important;
            background-color: white !important;
            color: #334155 !important;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05) !important;
        }
        .stTextInput input::placeholder {
            color: #94a3b8 !important;
            font-weight: 600 !important;
        }
        .stTextInput input:focus {
            border-color: #a855f7 !important;
            box-shadow: 0 0 0 4px rgba(168, 85, 247, 0.15) !important;
        }

        div[data-testid="stFormSubmitButton"] {
            display: flex;
            justify-content: center;
            width: 100%;
            margin-top: 15px !important;
        }
        div[data-testid="stFormSubmitButton"] > button {
            width: 70% !important;
            background: linear-gradient(135deg, #a855f7, #8b5cf6) !important;
            color: white !important;
            border: none !important;
            border-radius: 30px !important;
            padding: 14px !important;
            font-size: 1.1rem !important;
            font-weight: 700 !important;
            transition: all 0.3s ease;
            box-shadow: 0 10px 15px -3px rgba(139, 92, 246, 0.4) !important;
        }
        div[data-testid="stFormSubmitButton"] > button:hover {
            transform: translateY(-2px);
            box-shadow: 0 15px 20px -3px rgba(139, 92, 246, 0.5) !important;
        }

        .form-links {
            display: flex;
            justify-content: space-between;
            margin-top: 25px;
            padding: 0 15px;
        }
        .form-links a {
            color: #64748b;
            font-weight: 600;
            font-size: 0.9rem;
            text-decoration: none;
            transition: color 0.2s;
        }
        .form-links a:hover { color: #8b5cf6; }

        .powered-by {
            text-align: center;
            color: rgba(255, 255, 255, 0.9);
            font-size: 0.95rem;
            font-weight: 500;
            margin-top: 30px;
            text-shadow: 0 2px 4px rgba(0,0,0,0.6); 
            letter-spacing: 0.5px;
        }
        </style>
    """, unsafe_allow_html=True)

    c1, c2, c3 = st.columns([1, 2, 1])
    
    with c2:
        with st.form("login_form"):
            st.markdown('<div class="main-title">Sign In</div>', unsafe_allow_html=True)
            
            username = st.text_input("Username", placeholder="username or email", label_visibility="collapsed")
            st.write("") 
            password = st.text_input("Password", type="password", placeholder="password", label_visibility="collapsed")
            
            submit = st.form_submit_button("SIGN IN")
            
            st.markdown("""
                <div class="form-links">
                    <a href="#">Forgot password?</a>
                    <a href="#">Sign Up</a>
                </div>
            """, unsafe_allow_html=True)
            
            if submit:
                if username == "slogsolutions" and password == "slog2026":
                    st.session_state.logged_in = True
                    st.rerun()
                else:
                    st.error("Invalid credentials")

        st.markdown('<div class="powered-by">Powered by SLOG Solutions</div>', unsafe_allow_html=True)

else:
    # --- 5. MAIN APP UI ---
    st.set_page_config(page_title="Question Gen AI", layout="wide")

    with st.sidebar:
        if st.button("Logout"):
            st.session_state.logged_in = False
            st.rerun()
        st.divider()
        st.header("📌 Paper Details")
        title = st.text_input("Title", "")
        branch = st.text_input("Branch Name", "")
        b_code = st.text_input("Branch Code", "")
        sem = st.text_input("Semester", "")
        sub_name = st.text_input("Subject Name", "")
        sub_code = st.text_input("Subject Code", "")
        duration = st.text_input("Duration(Hrs)", "")

    st.title("🎓 Question Paper & Answer Key Generator")
    st.subheader("Powered by SLOG Solutions")

    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("### Section A/ भाग - क")
        n_a = st.number_input("Questions to Attempt", 1, 30, 10, key="n_a_input")
        m_a = st.number_input("Section Marks", 1, 100, 10, key="m_a_input")
        
        # ---  UI FOR SECTION A BREAKDOWN ---
        st.markdown("##### ⚙️ Question Breakdown")
        num_mcq = st.number_input("Number of MCQs", 0, 30, 4, key="mcq_input")
        num_fill = st.number_input("Number of Fill-ups", 0, 30, 4, key="fill_input")
        num_tf = st.number_input("Number of True/False", 0, 30, 4, key="tf_input")
        
        # Calculate Total automatically
        total_a = num_mcq + num_fill + num_tf
        st.info(f"**Total Section A Questions:** {total_a}")
        
        diff_a = st.selectbox("Difficulty", ["Easy", "Medium", "Hard"], key="d_a_input")

    with col2:
        st.markdown("### Section B/ भाग - ख")
        n_b = st.number_input("Questions to Attempt", 1, 30, 5, key="n_b_input")
        m_b = st.number_input("Section Marks", 1, 100, 15, key="m_b_input")
        total_b = st.number_input("Total Questions", n_b, n_b+10, n_b+2, key="t_b_input")
        diff_b = st.selectbox("Difficulty", ["Easy", "Medium", "Hard"], index=1, key="d_b_input")

    with col3:
        st.markdown("### Section C/ भाग - ग")
        n_c = st.number_input("Questions to Attempt", 1, 30, 5, key="n_c_input")
        m_c = st.number_input("Section Marks", 1, 100, 25, key="m_c_input")
        total_c = st.number_input("Total Questions", n_c, n_c+10, n_c+2, key="t_c_input")
        diff_c = st.selectbox("Difficulty", ["Easy", "Medium", "Hard"], index=2, key="d_c_input")

    st.divider()
    uploaded_file = st.file_uploader("Upload Study Material (PDF)", type="pdf")

    if st.button("🚀 Generate Question Paper & Answer Key") and uploaded_file:
        with st.status("AI is working...", expanded=True) as status:
            st.write("Reading PDF content...")
            raw_text = extract_text_from_pdf(uploaded_file)
            
            config = {
                "title": title, "branch_name": branch, "branch_code": b_code,
                "sem": sem, "subject_name": sub_name, "subject_code": sub_code,
                "duration": duration, "total_marks": m_a + m_b + m_c,
                
                # We now pass the specific breakdown alongside the total
                "n_a": n_a, "m_a": m_a, "total_a": total_a, "diff_a": diff_a,
                "num_mcq": num_mcq, "num_fill": num_fill, "num_tf": num_tf,
                
                "n_b": n_b, "m_b": m_b, "total_b": total_b, "diff_b": diff_b,
                "n_c": n_c, "m_c": m_c, "total_c": total_c, "diff_c": diff_c
            }
            
            try:
                data = get_llm_response(raw_text, config)
                st.write("Creating Word documents...")
                q_doc, a_doc = create_word_files(data, config)
                
                q_buf, a_buf = io.BytesIO(), io.BytesIO()
                q_doc.save(q_buf)
                a_doc.save(a_buf)
                
                st.session_state.qp_bytes = q_buf.getvalue()
                st.session_state.ak_bytes = a_buf.getvalue()
                st.session_state.current_sub_code = sub_code
                status.update(label="Generation Complete!", state="complete", expanded=False)
            except Exception as e:
                st.error(f"Error: {str(e)}")

    if st.session_state.qp_bytes is not None:
        st.divider()
        c1, c2 = st.columns(2)
        with c1:
            st.download_button("📥 Download Question Paper", st.session_state.qp_bytes, f"QP_{st.session_state.current_sub_code}.docx")
        with c2:
            st.download_button("📥 Download Answer Key", st.session_state.ak_bytes, f"AK_{st.session_state.current_sub_code}.docx")
